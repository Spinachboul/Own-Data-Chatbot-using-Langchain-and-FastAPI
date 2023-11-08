# Importing the libraries
from langchain.prompts import PromptTemplate
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.chains.question_answering import load_qa_chain
from langchain.llms import AzureOpenAI

from fastapi import FastAPI, File, UploadFile, Form, Header
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware

from typing import List, Optional, Union
import os
from dotenv import load_dotenv
import uvicorn
import io
import requests

from PyPDF2 import PdfReader
import pyxlsb
import tabula
import pandas as pd
from docx import Document
import markdown
import difflib
import json
import re

prompt_template = """Use the following pieces of context to answer the question at the end. 
Provide answer from input source documents / knowledge base / context only.
If the answer is found from the document, then please provide the reference page number.

 

{context}

 

Question: {question}
"""

 

PROMPT = PromptTemplate(
	template=prompt_template, input_variables=["context", "question"]
)

# Loading the environment
load_dotenv()

# Create a dictionary to store the document names and their text content
uploaded_documents = {}
questions : List[str]

# fastapi object creation
app = FastAPI()
app.add_middleware(
    CORSMiddleware, allow_origins=["*"], allow_headers=["*"], allow_methods=["*"], allow_credentials=True
)



# Mounting the static folder
app.mount("/static", StaticFiles(directory="static"), name="static")

# Initialize the knowledge_base as None
knowledge_base = None



# "Upload" Endpoint
@app.post("/upload")
async def upload_files(files: List[UploadFile]):
    """
    Extract the text from uploaded files and create a knowledge base.
    """
    global knowledge_base
    text = ""
    for file in files:
        # Appending text from all kinds of file formats: [.pdf, .txt, .docx, .csv, .xlsx, .xlsb]
        if file.content_type == "application/pdf":
            # Read PDF contents as binary
            contents = file.file.read()
            pdf_reader = PdfReader(io.BytesIO(contents))
            page_texts = []
            for i, page in enumerate(pdf_reader.pages):
                page_texts.append((i + 1, page.extract_text()))  # Store page number and text
            text += "".join(text for _, text in page_texts)
            uploaded_documents[file.filename] = page_texts  


        elif file.content_type == "text/plain":
            # Read plain text contents as string
            text += file.file.read().decode()
            uploaded_documents[file.filename].append((1, text))

        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Read DOCX contents as binary
            contents = file.file.read()
            doc_reader = Document(io.BytesIO(contents))
            doc_text = ""
            for paragraph in doc_reader.paragraphs:
                doc_text += paragraph.text + "\n"
            text += doc_text
            uploaded_documents[file.filename].append((1, text))

        elif file.content_type == "application/msword":
            # Read DOC contents as binary
            contents = file.file.read()
            doc_reader = Document(io.BytesIO(contents))
            doc_text = ""
            for paragraph in doc_reader.paragraphs:
                doc_text += paragraph.text + "\n"
            text += doc_text
            uploaded_documents[file.filename].append((1, text))

        elif file.content_type == "text/csv":
            # Read CSV contents as string
            csv_text = pd.read_csv(file.file)
            text += csv_text.to_string(index=False)
            uploaded_documents[file.filename].append((1, text))

        elif file.content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            # Read XLSX contents as binary
            contents = file.file.read()
            xls = pd.ExcelFile(io.BytesIO(contents))
            for sheet_name in xls.sheet_names:
                sheet_text = xls.parse(sheet_name, engine='GenpactGpt')
                text += sheet_text.to_csv(index=True)
            uploaded_documents[file.filename].append((1, text))

        elif file.content_type == "application/vnd.ms-excel.sheet.binary.macroEnabled.12":
            # Read XLSB contents as binary
            contents = file.file.read()
            xlsb_reader = pyxlsb.open_workbook(io.BytesIO(contents))
            for sheet_name in xlsb_reader.sheets:
                sheet_text = ""
                for row in xlsb_reader.get_sheet(sheet_name):
                    row_text = "\t".join([str(cell.v) for cell in row])
                    sheet_text += row_text + "\n"
                text += sheet_text
            uploaded_documents[file.filename].append((1, text))
    # Read and extract questions from the uploaded CSV file
    


    # Creating a texx-splitter object
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=6000,
        chunk_overlap=600,
        length_function=len
    )

    # Splitting the text into chunks
    chunks = text_splitter.split_text(text)

    # Creating embeddings for the OpenAI Model
    embeddings = OpenAIEmbeddings(
        engine="text-embedding-ada-002",
        deployment_id="text-embedding-ada-002",
        chunk_size=1
    )
    knowledge_base = Chroma.from_texts(chunks, embeddings)

    return {"status": "success", "message": "Knowledge base created successfully."}

# "Question-Answering" Endpoint
@app.post("/qa")
async def perform_qa(query: str = Form(...), doctype: str = Form(...), category_value: str = Form(...)):
    global knowledge_base
    global uploaded_documents

    if knowledge_base is None:
        return {
            "status": "error",
            "message": "No knowledge base available. Please upload a file first."
        }

    if doctype == "financeFunctional":
        prompt_file = 'main.csv'
        df = pd.read_csv(prompt_file)
        df = df[df['CATEGORY'] == category_value]
        questions = df.iloc[:, 2].tolist()
        field_name = df.iloc[:, 1].tolist()

        answers = []
        for i in range(len(questions)):
            prompt = questions[i]
            fName = field_name[i]
            try:
                docs = knowledge_base.similarity_search(prompt)
                llm = AzureOpenAI(
                    deployment_name="Genpacttest",
                    model_name="gpt-35-turbo",
                    temperature=0.5
                )
                chain = load_qa_chain(llm, chain_type="stuff")
                response = chain.run(input_documents=docs, question=prompt)
                filter_response = response.split("\n\n")[0].strip('\'"').replace(':', '').replace(',', '').replace('`', '')
            except:
                continue

            max_similarity = 0.0
            page_number = None
            document_name = None

            for doc_name, doc_pages in uploaded_documents.items():
                for page_num, doc_text in doc_pages:
                    similarity = difflib.SequenceMatcher(None, doc_text, filter_response).ratio()
                    if similarity >= max_similarity and filter_response != "I don't know":
                        max_similarity = similarity
                        page_number = page_num
                        document_name = doc_name

            answers.append({
                "Field_Name": fName,
                "Answer": filter_response,
                "Page Number": page_number,
                "Document Name": document_name
            })

        return {"Answers": answers}
    else:
        if query is None:
            return {
                "status": "error",
                "message": "No query provided."
            }

        docs = knowledge_base.similarity_search(query)
        llm = AzureOpenAI(
            deployment_name="Genpacttest",
            model_name="gpt-35-turbo",
            temperature=0.6
        )
        chain = load_qa_chain(llm, chain_type="stuff", prompt=PROMPT)
        response = chain.run(input_documents=docs, question=query)
        filtered_response = response.split("\n\n")[0].strip('\'"').replace(':', '').replace(',', '').replace('`', '')

        return {"answer": filtered_response}

In this code, also return the page Number and document name in the else part, when the financeFunctional is not selected as an option
